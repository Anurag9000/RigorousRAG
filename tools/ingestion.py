"""Safe document parsing, optional OCR, redaction, and semantic sections."""

from __future__ import annotations

import hashlib
import io
import mimetypes
import os
import re
import statistics
import uuid
import zipfile
from pathlib import Path, PurePosixPath
from typing import Any, Dict, List, Optional, Sequence, Tuple

import docx
import fitz

from tools.ingestion_models import DocumentSection, IngestedDocument, IngestionResult
from tools.security import DEFAULT_MAX_UPLOAD_BYTES, normalize_owner_id

try:
    import pytesseract
    from PIL import Image
except ImportError:  # OCR is an optional runtime capability.
    pytesseract = None  # type: ignore[assignment]
    Image = None  # type: ignore[assignment]

_ALLOWED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}
_PDF_MIME = "application/pdf"
_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_ENABLE_OCR = os.getenv("ENABLE_OCR", "false").lower() in {"1", "true", "yes"}
_OCR_MAX_PAGES = max(1, min(int(os.getenv("OCR_MAX_PAGES", "50")), 500))
_OCR_DPI = max(100, min(int(os.getenv("OCR_DPI", "200")), 400))
_OCR_TIMEOUT_SECONDS = max(1, min(int(os.getenv("OCR_TIMEOUT_SECONDS", "30")), 300))
_OCR_MIN_TEXT_CHARS = max(0, min(int(os.getenv("OCR_MIN_TEXT_CHARS", "40")), 2000))
_MAX_PDF_PAGES = max(1, min(int(os.getenv("MAX_PDF_PAGES", "2000")), 10_000))
_MAX_PDF_RENDER_PIXELS = max(
    1_000_000,
    min(int(os.getenv("MAX_PDF_RENDER_PIXELS", "40000000")), 250_000_000),
)
_MAX_EXTRACTED_CHARS = max(
    100_000,
    min(int(os.getenv("MAX_EXTRACTED_CHARS", "5000000")), 50_000_000),
)
_MAX_DOCX_MEMBERS = max(10, min(int(os.getenv("MAX_DOCX_MEMBERS", "10000")), 100_000))
_MAX_DOCX_UNCOMPRESSED_BYTES = max(
    DEFAULT_MAX_UPLOAD_BYTES,
    min(
        int(os.getenv("MAX_DOCX_UNCOMPRESSED_BYTES", "200000000")),
        2_000_000_000,
    ),
)
_MAX_DOCX_COMPRESSION_RATIO = max(
    10.0,
    min(float(os.getenv("MAX_DOCX_COMPRESSION_RATIO", "1000")), 100_000.0),
)
_EMAIL_RE = re.compile(r"(?<![\w.+-])[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}(?![\w.-])")
_PHONE_RE = re.compile(r"(?<!\w)(?:\+?\d{1,3}[\s().-]*)?(?:\d[\s().-]*){7,14}\d(?!\w)")
_ADDRESS_RE = re.compile(
    r"\b\d{1,6}\s+[A-Za-z0-9.' -]{2,80}\s+"
    r"(?:Street|St|Avenue|Ave|Road|Rd|Lane|Ln|Drive|Dr|Boulevard|Blvd|"
    r"Way|Court|Ct|Place|Pl|Highway|Hwy)\b\.?",
    flags=re.IGNORECASE,
)
_IPV4_RE = re.compile(
    r"\b(?:(?:25[0-5]|2[0-4]\d|1?\d?\d)\.){3}"
    r"(?:25[0-5]|2[0-4]\d|1?\d?\d)\b"
)
_CARD_RE = re.compile(r"(?<!\d)(?:\d[ -]*?){13,19}(?!\d)")
_DOI_RE = re.compile(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+\b", re.IGNORECASE)
_YEAR_RE = re.compile(r"\b(?:19|20)\d{2}\b")
_AUTHOR_LINE_RE = re.compile(r"^(?:authors?|by)\s*[:—-]\s*(.+)$", re.IGNORECASE)
_SENTENCE_BOUNDARY_RE = re.compile(r"(?<=[.!?])\s+")


class OCRUnavailableError(RuntimeError):
    """Raised when OCR was explicitly enabled but its runtime is unavailable."""


class IngestionLimitError(ValueError):
    """Raised when a document exceeds a configured parser-complexity budget."""


def detect_mime_type(file_path: str) -> str:
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        return _PDF_MIME
    if suffix == ".docx":
        return _DOCX_MIME
    if suffix == ".md":
        return "text/markdown"
    if suffix == ".txt":
        return "text/plain"
    return mimetypes.guess_type(file_path)[0] or "application/octet-stream"


def _luhn_valid(candidate: str) -> bool:
    digits = [int(value) for value in re.sub(r"\D", "", candidate)]
    if not 13 <= len(digits) <= 19 or len(set(digits)) == 1:
        return False
    total = 0
    parity = len(digits) % 2
    for index, digit in enumerate(digits):
        if index % 2 == parity:
            digit *= 2
            if digit > 9:
                digit -= 9
        total += digit
    return total % 10 == 0


def redact_text(text: str) -> str:
    """Best-effort masking; this is not a proof of anonymization."""

    value = text or ""
    value = _EMAIL_RE.sub("[REDACTED_EMAIL]", value)
    value = _ADDRESS_RE.sub("[REDACTED_ADDRESS]", value)
    value = _IPV4_RE.sub("[REDACTED_IP]", value)

    def redact_card(match: re.Match[str]) -> str:
        raw = match.group(0)
        return "[REDACTED_PAYMENT_CARD]" if _luhn_valid(raw) else raw

    value = _CARD_RE.sub(redact_card, value)
    value = _PHONE_RE.sub("[REDACTED_PHONE]", value)
    return value


def extract_academic_metadata(text: str) -> Dict[str, Any]:
    sample = (text or "")[:8000]
    metadata: Dict[str, Any] = {}
    doi = _DOI_RE.search(sample)
    if doi:
        metadata["doi"] = doi.group(0).rstrip(".,;)")
    years = _YEAR_RE.findall(sample)
    if years:
        metadata["year"] = years[0]
    lines = [line.strip() for line in sample.splitlines() if line.strip()]
    title_candidates: List[str] = []
    for line in lines[:20]:
        lowered = line.lower().strip(":")
        if lowered in {"abstract", "introduction", "methods", "results", "references"}:
            continue
        if _DOI_RE.search(line) or len(line) < 5 or len(line) > 300:
            continue
        if _AUTHOR_LINE_RE.match(line):
            continue
        title_candidates.append(line)
    if title_candidates:
        metadata["extracted_title"] = title_candidates[0]
    for line in lines[:30]:
        author_match = _AUTHOR_LINE_RE.match(line)
        if author_match:
            metadata["authors"] = author_match.group(1)[:1000]
            break
    return metadata


def _split_oversized_unit(unit: str, max_chars: int) -> List[str]:
    unit = unit.strip()
    if not unit:
        return []
    if len(unit) <= max_chars:
        return [unit]
    sentences = [item.strip() for item in _SENTENCE_BOUNDARY_RE.split(unit) if item.strip()]
    if len(sentences) <= 1:
        sentences = []
    result: List[str] = []
    current = ""
    for sentence in sentences:
        if len(sentence) > max_chars:
            if current:
                result.append(current)
                current = ""
            word_buffer = ""
            for word in sentence.split():
                candidate = f"{word_buffer} {word}".strip()
                if word_buffer and len(candidate) > max_chars:
                    result.append(word_buffer)
                    word_buffer = word
                else:
                    word_buffer = candidate
            if word_buffer:
                result.append(word_buffer)
            continue
        candidate = f"{current} {sentence}".strip()
        if current and len(candidate) > max_chars:
            result.append(current)
            current = sentence
        else:
            current = candidate
    if current:
        result.append(current)
    if result:
        return result
    return [unit[index:index + max_chars] for index in range(0, len(unit), max_chars)]


def _chunk_text_semantically(text: str, max_chars: int = 1500) -> List[str]:
    if max_chars < 100:
        max_chars = max(1, max_chars)
    paragraphs = [item.strip() for item in re.split(r"\n\s*\n", text or "") if item.strip()]
    chunks: List[str] = []
    current = ""
    for paragraph in paragraphs:
        for unit in _split_oversized_unit(paragraph, max_chars):
            candidate = f"{current}\n\n{unit}".strip()
            if current and len(candidate) > max_chars:
                chunks.append(current)
                current = unit
            else:
                current = candidate
    if current:
        chunks.append(current)
    return chunks


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _validate_docx_archive(path: Path) -> None:
    try:
        with zipfile.ZipFile(path) as archive:
            infos = archive.infolist()
            if len(infos) > _MAX_DOCX_MEMBERS:
                raise IngestionLimitError(
                    f"The DOCX contains more than {_MAX_DOCX_MEMBERS} archive members."
                )
            names = [info.filename for info in infos]
            if len(names) != len(set(names)):
                raise ValueError("The DOCX package contains duplicate archive members.")
            required = {"[Content_Types].xml", "word/document.xml"}
            if not required.issubset(set(names)):
                raise ValueError("The DOCX package is missing required document members.")
            total_uncompressed = 0
            total_compressed = 0
            for info in infos:
                member = PurePosixPath(info.filename)
                if member.is_absolute() or ".." in member.parts:
                    raise ValueError("The DOCX package contains an unsafe archive path.")
                mode = (info.external_attr >> 16) & 0o170000
                if mode == 0o120000:
                    raise ValueError("The DOCX package contains a symbolic-link member.")
                if info.flag_bits & 0x1:
                    raise ValueError("Encrypted DOCX archive members are not supported.")
                total_uncompressed += max(0, int(info.file_size))
                total_compressed += max(0, int(info.compress_size))
                if total_uncompressed > _MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise IngestionLimitError(
                        "The DOCX uncompressed content exceeds the configured limit."
                    )
            ratio = total_uncompressed / max(total_compressed, 1)
            if ratio > _MAX_DOCX_COMPRESSION_RATIO:
                raise IngestionLimitError(
                    "The DOCX compression ratio exceeds the configured safety limit."
                )
    except zipfile.BadZipFile as exc:
        raise ValueError("The DOCX package is malformed.") from exc


def _validate_file(path: Path) -> str:
    if path.is_symlink():
        raise ValueError("Symbolic-link input files are not supported.")
    if not path.exists() or not path.is_file():
        raise ValueError("The requested input file does not exist.")
    size = path.stat().st_size
    if size <= 0:
        raise ValueError("The input file is empty.")
    if size > DEFAULT_MAX_UPLOAD_BYTES:
        raise ValueError(
            f"The input file exceeds the {DEFAULT_MAX_UPLOAD_BYTES}-byte upload limit."
        )
    suffix = path.suffix.lower()
    if suffix not in _ALLOWED_SUFFIXES:
        raise ValueError(
            f"Unsupported file type '{suffix or 'none'}'. "
            f"Allowed: {', '.join(sorted(_ALLOWED_SUFFIXES))}."
        )
    with path.open("rb") as handle:
        prefix = handle.read(4096)
    if suffix == ".pdf" and not prefix.startswith(b"%PDF-"):
        raise ValueError("The file extension is .pdf but the content is not a PDF.")
    if suffix == ".docx":
        if not prefix.startswith(b"PK"):
            raise ValueError("The file extension is .docx but the content is not a ZIP package.")
        _validate_docx_archive(path)
    if suffix in {".txt", ".md"} and b"\x00" in prefix:
        raise ValueError("The text file appears to contain binary data.")
    return detect_mime_type(str(path))


def _extract_tables_from_page(page: fitz.Page) -> List[str]:
    try:
        finder = page.find_tables()
    except Exception:
        return []
    rendered: List[str] = []
    for table in getattr(finder, "tables", []) or []:
        try:
            rows = table.extract()
        except Exception:
            continue
        lines = [
            "\t".join("" if cell is None else str(cell).strip() for cell in row)
            for row in rows
        ]
        text = "\n".join(line for line in lines if line.strip())
        if text:
            rendered.append(text)
    return rendered


def _page_heading_candidates(page: fitz.Page) -> List[Tuple[str, float, bool]]:
    try:
        blocks = page.get_text("dict", sort=True).get("blocks", [])
    except Exception:
        return []
    candidates: List[Tuple[str, float, bool]] = []
    for block in blocks:
        for line in block.get("lines", []) or []:
            spans = line.get("spans", []) or []
            text = " ".join(
                str(span.get("text") or "").strip() for span in spans
            ).strip()
            if not text or len(text) > 250:
                continue
            sizes = [
                float(span.get("size") or 0.0)
                for span in spans
                if span.get("text")
            ]
            size = max(sizes) if sizes else 0.0
            bold = any(int(span.get("flags") or 0) & 16 for span in spans)
            candidates.append((text, size, bold))
    return candidates


def _ocr_page(page: fitz.Page, page_number: int) -> str:
    if not _ENABLE_OCR:
        return ""
    if pytesseract is None or Image is None:
        raise OCRUnavailableError(
            "OCR is enabled but Pillow/pytesseract is unavailable. Install OCR dependencies."
        )
    width_pixels = max(1.0, float(page.rect.width) * _OCR_DPI / 72.0)
    height_pixels = max(1.0, float(page.rect.height) * _OCR_DPI / 72.0)
    if width_pixels * height_pixels > _MAX_PDF_RENDER_PIXELS:
        raise IngestionLimitError(
            f"PDF page {page_number} exceeds the configured OCR render-pixel limit."
        )
    pixmap = page.get_pixmap(dpi=_OCR_DPI, alpha=False)
    try:
        with Image.open(io.BytesIO(pixmap.tobytes("png"))) as image:
            text = pytesseract.image_to_string(
                image,
                config="--psm 3",
                timeout=_OCR_TIMEOUT_SECONDS,
            )
    except Exception as exc:
        not_found_type = getattr(pytesseract, "TesseractNotFoundError", ())
        if not_found_type and isinstance(exc, not_found_type):
            raise OCRUnavailableError("The Tesseract executable is unavailable.") from exc
        raise
    return "\n".join(line.rstrip() for line in text.splitlines()).strip()


def _extract_sections_from_pdf(
    document: fitz.Document,
) -> Tuple[List[DocumentSection], Dict[str, List[int]]]:
    if len(document) > _MAX_PDF_PAGES:
        raise IngestionLimitError(
            f"The PDF contains more than the configured {_MAX_PDF_PAGES}-page limit."
        )
    page_payloads: List[Tuple[int, str, List[Tuple[str, float, bool]], bool]] = []
    all_sizes: List[float] = []
    extracted_chars = 0
    ocr_stats: Dict[str, List[int]] = {
        "attempted": [],
        "successful": [],
        "empty": [],
        "failed": [],
        "skipped_limit": [],
    }
    for page_index, page in enumerate(document):
        page_number = page_index + 1
        text = page.get_text("text", sort=True).strip()
        used_ocr = False
        low_text = len(re.sub(r"\s+", "", text)) < _OCR_MIN_TEXT_CHARS
        if low_text and _ENABLE_OCR:
            if len(ocr_stats["attempted"]) >= _OCR_MAX_PAGES:
                ocr_stats["skipped_limit"].append(page_number)
            else:
                ocr_stats["attempted"].append(page_number)
                try:
                    ocr_text = _ocr_page(page, page_number)
                except OCRUnavailableError:
                    raise
                except Exception:
                    ocr_stats["failed"].append(page_number)
                    ocr_text = ""
                if ocr_text:
                    text = ocr_text
                    used_ocr = True
                    ocr_stats["successful"].append(page_number)
                elif page_number not in ocr_stats["failed"]:
                    ocr_stats["empty"].append(page_number)
        tables = _extract_tables_from_page(page) if not used_ocr else []
        if tables:
            text = f"{text}\n\n" + "\n\n".join(f"[TABLE]\n{table}" for table in tables)
        extracted_chars += len(text)
        if extracted_chars > _MAX_EXTRACTED_CHARS:
            raise IngestionLimitError(
                "PDF extraction exceeded the configured text-character limit."
            )
        candidates = _page_heading_candidates(page) if not used_ocr else []
        all_sizes.extend(size for _, size, _ in candidates if size > 0)
        page_payloads.append((page_number, text.strip(), candidates, used_ocr))
    median_size = statistics.median(all_sizes) if all_sizes else 0.0
    sections: List[DocumentSection] = []
    for page_number, page_text, candidates, used_ocr in page_payloads:
        if not page_text:
            continue
        heading = f"Page {page_number}"
        if used_ocr:
            heading += " (OCR)"
        else:
            for text, size, bold in candidates:
                if median_size and (
                    (bold and size >= median_size) or size >= median_size * 1.25
                ) and 3 <= len(text) <= 180:
                    heading = text
                    break
        sections.append(
            DocumentSection(
                title=heading[:500],
                content=page_text,
                page_number=page_number,
            )
        )
    return sections, ocr_stats


def _csv_pages(values: Sequence[int]) -> str:
    return ",".join(str(value) for value in values)


def _ingest_pdf(path: Path) -> IngestionResult:
    try:
        document = fitz.open(path)
    except Exception as exc:
        return IngestionResult(success=False, error=f"Could not open PDF: {type(exc).__name__}.")
    try:
        if document.needs_pass:
            return IngestionResult(success=False, error="Encrypted PDFs are not supported.")
        sections, ocr_stats = _extract_sections_from_pdf(document)
        text = "\n\n".join(section.content for section in sections).strip()
        if not text:
            if _ENABLE_OCR and (pytesseract is None or Image is None):
                error = "The PDF requires OCR, but OCR dependencies are unavailable."
            elif _ENABLE_OCR and ocr_stats["skipped_limit"]:
                error = (
                    "No text was extracted within the configured OCR-attempt limit "
                    f"({_OCR_MAX_PAGES})."
                )
            elif _ENABLE_OCR and ocr_stats["failed"]:
                error = "OCR failed on every attempted low-text page."
            elif _ENABLE_OCR:
                error = "OCR completed but produced no indexable text."
            else:
                error = "The PDF contains no extractable text; set ENABLE_OCR=true."
            return IngestionResult(success=False, error=error)
        metadata = dict(document.metadata or {})
        title = str(metadata.get("title") or "").strip() or None
        safe_metadata = {
            key: value
            for key, value in metadata.items()
            if isinstance(value, (str, int, float, bool)) and value not in (None, "")
        }
        warnings: List[str] = []
        if ocr_stats["empty"]:
            warnings.append("OCR returned no text for some attempted pages.")
        if ocr_stats["failed"]:
            warnings.append("OCR failed or timed out for some attempted pages.")
        if ocr_stats["skipped_limit"]:
            warnings.append("Some low-text pages were skipped after the OCR attempt limit.")
        safe_metadata.update(
            {
                "page_count": len(document),
                "ocr_enabled": _ENABLE_OCR,
                "ocr_attempted_count": len(ocr_stats["attempted"]),
                "ocr_attempted_pages": _csv_pages(ocr_stats["attempted"]),
                "ocr_page_count": len(ocr_stats["successful"]),
                "ocr_pages": _csv_pages(ocr_stats["successful"]),
                "ocr_empty_pages": _csv_pages(ocr_stats["empty"]),
                "ocr_failed_pages": _csv_pages(ocr_stats["failed"]),
                "ocr_skipped_limit_pages": _csv_pages(ocr_stats["skipped_limit"]),
                "extraction_warnings": " ".join(warnings),
            }
        )
        return IngestionResult(
            success=True,
            document=IngestedDocument(
                id="pending",
                filename=path.name,
                file_path=str(path),
                mime_type=_PDF_MIME,
                title=title,
                text=text,
                sections=sections,
                metadata=safe_metadata,
            ),
        )
    except (OCRUnavailableError, IngestionLimitError) as exc:
        return IngestionResult(success=False, error=str(exc))
    except Exception as exc:
        return IngestionResult(
            success=False,
            error=f"PDF extraction failed: {type(exc).__name__}.",
        )
    finally:
        document.close()


def _ingest_docx(path: Path) -> IngestionResult:
    try:
        document = docx.Document(str(path))
        sections: List[DocumentSection] = []
        current_title = "Document"
        current_lines: List[str] = []
        extracted_chars = 0

        def add_chars(value: str) -> None:
            nonlocal extracted_chars
            extracted_chars += len(value)
            if extracted_chars > _MAX_EXTRACTED_CHARS:
                raise IngestionLimitError(
                    "DOCX extraction exceeded the configured text-character limit."
                )

        def flush() -> None:
            nonlocal current_lines
            content = "\n".join(line for line in current_lines if line.strip()).strip()
            if content:
                sections.append(DocumentSection(title=current_title[:500], content=content))
            current_lines = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            add_chars(text)
            style_name = str(getattr(paragraph.style, "name", "") or "")
            if style_name.lower().startswith("heading"):
                flush()
                current_title = text
            else:
                current_lines.append(text)
        flush()
        for table_index, table in enumerate(document.tables, start=1):
            rows = ["\t".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            table_text = "\n".join(row for row in rows if row.strip()).strip()
            if table_text:
                add_chars(table_text)
                sections.append(
                    DocumentSection(title=f"Table {table_index}", content=table_text)
                )
        text = "\n\n".join(section.content for section in sections).strip()
        if not text:
            return IngestionResult(success=False, error="The DOCX contains no extractable text.")
        properties = document.core_properties
        metadata: Dict[str, Any] = {}
        for key in ("author", "subject", "keywords", "comments", "category"):
            value = getattr(properties, key, None)
            if value:
                metadata[key] = str(value)
        title = str(getattr(properties, "title", "") or "").strip() or None
        return IngestionResult(
            success=True,
            document=IngestedDocument(
                id="pending",
                filename=path.name,
                file_path=str(path),
                mime_type=_DOCX_MIME,
                title=title,
                text=text,
                sections=sections,
                metadata=metadata,
            ),
        )
    except IngestionLimitError as exc:
        return IngestionResult(success=False, error=str(exc))
    except Exception as exc:
        return IngestionResult(
            success=False,
            error=f"Could not parse DOCX: {type(exc).__name__}.",
        )


def _ingest_text(path: Path, mime_type: str) -> IngestionResult:
    try:
        raw = path.read_bytes()
    except OSError as exc:
        return IngestionResult(success=False, error=f"Could not read text file: {type(exc).__name__}.")
    decoded: Optional[str] = None
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            decoded = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if decoded is None:
        return IngestionResult(success=False, error="The text encoding could not be decoded.")
    text = decoded.strip()
    if not text:
        return IngestionResult(success=False, error="The text document is empty.")
    if len(text) > _MAX_EXTRACTED_CHARS:
        return IngestionResult(
            success=False,
            error="Text extraction exceeded the configured character limit.",
        )
    return IngestionResult(
        success=True,
        document=IngestedDocument(
            id="pending",
            filename=path.name,
            file_path=str(path),
            mime_type=mime_type,
            title=None,
            text=text,
            sections=[DocumentSection(title="Full Text", content=text)],
            metadata={},
        ),
    )


def _redact_sections(sections: Sequence[DocumentSection]) -> List[DocumentSection]:
    redacted: List[DocumentSection] = []
    for section in sections:
        content = redact_text(section.content).strip()
        if not content:
            continue
        chunks = _chunk_text_semantically(content, max_chars=6000) or [content]
        for chunk_index, chunk in enumerate(chunks):
            title = section.title
            if len(chunks) > 1:
                title = f"{title} — Part {chunk_index + 1}"
            redacted.append(
                DocumentSection(
                    title=title[:500],
                    content=chunk,
                    page_number=section.page_number,
                )
            )
    return redacted


def ingest_file(file_path: str, owner_id: str = "default_user") -> IngestionResult:
    path = Path(file_path)
    try:
        owner = normalize_owner_id(owner_id)
        initial_stat = path.stat()
        mime_type = _validate_file(path)
        source_hash = _sha256_file(path)
    except Exception as exc:
        return IngestionResult(success=False, error=str(exc))
    try:
        if path.suffix.lower() == ".pdf":
            result = _ingest_pdf(path)
        elif path.suffix.lower() == ".docx":
            result = _ingest_docx(path)
        else:
            result = _ingest_text(path, mime_type)
    except Exception as exc:
        return IngestionResult(
            success=False,
            error=f"Document parsing failed: {type(exc).__name__}.",
        )
    if not result.success or not result.document:
        return result

    try:
        final_stat = path.stat()
    except OSError:
        return IngestionResult(success=False, error="The input file disappeared during parsing.")
    if (
        initial_stat.st_size != final_stat.st_size
        or initial_stat.st_mtime_ns != final_stat.st_mtime_ns
        or getattr(initial_stat, "st_ino", None) != getattr(final_stat, "st_ino", None)
    ):
        return IngestionResult(success=False, error="The input file changed during parsing.")

    document = result.document
    redacted_text = redact_text(document.text).strip()
    redacted_sections = _redact_sections(document.sections)
    if not redacted_text or not redacted_sections:
        return IngestionResult(
            success=False,
            error="No indexable text remained after parsing.",
        )
    redacted_hash = hashlib.sha256(redacted_text.encode("utf-8")).hexdigest()
    stable_id = str(
        uuid.uuid5(uuid.NAMESPACE_URL, f"rigorousrag:{owner}:{source_hash}")
    )
    extracted = extract_academic_metadata(redacted_text)
    document.id = stable_id
    document.text = redacted_text
    document.sections = redacted_sections
    document.title = (
        document.title
        or extracted.get("extracted_title")
        or path.stem.replace("_", " ").replace("-", " ")
    )
    document.metadata.update(extracted)
    document.metadata.update(
        {
            "owner_id": owner,
            "content_sha256": redacted_hash,
            "file_size_bytes": final_stat.st_size,
            "redaction": "best_effort_regex_masking",
            "document_identity": "owner_and_source_sha256",
        }
    )
    return IngestionResult(success=True, document=document)
